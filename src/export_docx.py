\
from __future__ import annotations
from typing import List, Dict, Any
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ND30_MARGINS_CM = dict(top=2.0, bottom=2.0, left=3.0, right=2.0)
FONT_NAME = "Times New Roman"
FONT_SIZE = 13
LINE_SPACING = 1.15

def _apply_style(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(ND30_MARGINS_CM["top"])
    section.bottom_margin = Cm(ND30_MARGINS_CM["bottom"])
    section.left_margin = Cm(ND30_MARGINS_CM["left"])
    section.right_margin = Cm(ND30_MARGINS_CM["right"])

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

def _p(doc: Document, text: str, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    return p

def export_exam_docx(meta: Dict[str, Any], questions: List[Dict[str, Any]]) -> bytes:
    """
    Xuất 'Đề' + 'Đáp án/Hướng dẫn' đơn giản. Bỏ quốc hiệu-tiêu ngữ.
    """
    doc = Document()
    _apply_style(doc)

    _p(doc, (meta.get("title") or "ĐỀ KIỂM TRA").upper(), bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    info = f"Môn: {meta.get('subject','')}  •  Lớp: {meta.get('grade','5')}  •  Thời gian: {meta.get('time','40')} phút"
    _p(doc, info, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    _p(doc, "I. PHẦN CÂU HỎI", bold=True)

    for i, q in enumerate(questions, 1):
        _p(doc, f"Câu {i}. ({q.get('points',1)} điểm) {q.get('qtype','')} - {q.get('level','')}", bold=True)
        content = q.get("content", {})
        _p(doc, content.get("stem",""))

        qt = q.get("qtype","")
        if qt == "Trắc nghiệm nhiều lựa chọn":
            opts = content.get("options", {})
            for k in ["A","B","C","D"]:
                if k in opts:
                    _p(doc, f"{k}. {opts.get(k,'')}")
        elif qt == "Đúng/Sai":
            for j, it in enumerate(content.get("true_false", []), 1):
                _p(doc, f"{j}) {it.get('statement','')}")
        elif qt == "Nối cột":
            mt = content.get("matching", {})
            left = mt.get("left", [])
            right = mt.get("right", [])
            rows = max(len(left), len(right), 1) + 1
            t = doc.add_table(rows=rows, cols=2)
            t.style = "Table Grid"
            t.cell(0,0).text = "Cột A"
            t.cell(0,1).text = "Cột B"
            for r in range(rows-1):
                t.cell(r+1,0).text = left[r] if r < len(left) else ""
                t.cell(r+1,1).text = right[r] if r < len(right) else ""
        elif qt == "Điền khuyết":
            fb = content.get("fill_blank", {})
            _p(doc, fb.get("text",""))
        else:
            es = content.get("essay", {})
            _p(doc, es.get("prompt",""))

    # đáp án
    doc.add_page_break()
    _p(doc, "ĐÁP ÁN - HƯỚNG DẪN", bold=True)
    for i, q in enumerate(questions, 1):
        qt = q.get("qtype","")
        content = q.get("content", {})
        _p(doc, f"Câu {i}:", bold=True)
        if qt == "Trắc nghiệm nhiều lựa chọn":
            _p(doc, f"Đáp án: {content.get('correct_answer','')}")
        elif qt == "Đúng/Sai":
            tf = content.get("true_false", [])
            ans = ", ".join([f"{j+1}={'Đ' if it.get('answer') else 'S'}" for j, it in enumerate(tf)])
            _p(doc, f"Đáp án: {ans}")
        elif qt == "Nối cột":
            _p(doc, f"Đáp án: {mt.get('answer', {})}" if (mt := content.get("matching", {})) else "Đáp án: {}")
        elif qt == "Điền khuyết":
            _p(doc, f"Đáp án: {content.get('fill_blank', {}).get('answer','')}")
        else:
            rb = content.get("essay", {}).get("rubric", [])
            if rb:
                _p(doc, "Gợi ý chấm: " + "; ".join([str(x) for x in rb]))
            else:
                _p(doc, "Gợi ý chấm: (GV tự chấm theo đáp án/ý chính)")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
